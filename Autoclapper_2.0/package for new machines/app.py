from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
import os

app = Flask(__name__)
CORS(app)  # Allow requests from the HTML file

# HARDCODED TEMPLATE PATHS
SCRIPT_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'templates', 'script_template.docx')
CLAPPER_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'templates', 'clapper_blank.png')

# CLAPPER TEXT CONFIGURATION
CLAPPER_FONT_SIZE = 40  # Adjust this value to change text size
CLAPPER_OFFSET_X = 0    # Horizontal offset (negative = left, positive = right)
CLAPPER_OFFSET_Y = 0    # Vertical offset (negative = up, positive = down)

def generate_script_document(data, output_dir):
    """
    Generate the script document from template and user data.
    
    Template structure:
    - Table 1: Header (untouched)
    - Table 2: Job details (5 rows, 4 columns)
    - Table 3: Script content (2 rows, 2 columns)
    - Table 4: Generator signature
    """
    try:
        # Check if template exists
        if not os.path.exists(SCRIPT_TEMPLATE_PATH):
            return None, f"Script template not found at: {SCRIPT_TEMPLATE_PATH}"
        
        # Load the template
        doc = Document(SCRIPT_TEMPLATE_PATH)
        
        # Helper function to set cell text with formatting
        def set_cell_text(cell, text, bold=True, red=False):
            """Set text in a cell with Arial 11 formatting."""
            # Clear all existing paragraphs
            for paragraph in cell.paragraphs:
                paragraph.clear()
            
            # Use the first paragraph (or create if none exist)
            if not cell.paragraphs:
                paragraph = cell.add_paragraph()
            else:
                paragraph = cell.paragraphs[0]
            
            # Add the text run with formatting
            run = paragraph.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.bold = bold
            
            if red:
                run.font.color.rgb = RGBColor(178, 34, 34)  # Medium-dark red (Firebrick)
            
            # Remove any extra paragraphs that might cause page breaks
            while len(cell.paragraphs) > 1:
                cell._element.remove(cell.paragraphs[-1]._element)
        
        def set_cell_text_with_formatting(cell, text, bold=False, red=False):
            """Set text in a cell with Arial 11 formatting, handling formatting markers:
            /// = yellow highlight, ** = bold, * = italic, __ = underline"""
            from docx.enum.text import WD_COLOR_INDEX
            
            # Clear all existing paragraphs
            for paragraph in cell.paragraphs:
                paragraph.clear()
            
            # Use the first paragraph (or create if none exist)
            if not cell.paragraphs:
                paragraph = cell.add_paragraph()
            else:
                paragraph = cell.paragraphs[0]
            
            # Process the text character by character
            i = 0
            while i < len(text):
                # Check for highlight (///)
                if text[i:i+3] == '///':
                    end = text.find('///', i + 3)
                    if end != -1:
                        highlighted_text = text[i+3:end]
                        run = paragraph.add_run(highlighted_text)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.bold = True  # Highlighted text is always bold
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        i = end + 3
                        continue
                
                # Check for bold (**)
                elif text[i:i+2] == '**':
                    end = text.find('**', i + 2)
                    if end != -1:
                        bold_text = text[i+2:end]
                        run = paragraph.add_run(bold_text)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.bold = True
                        i = end + 2
                        continue
                
                # Check for underline (__)
                elif text[i:i+2] == '__':
                    end = text.find('__', i + 2)
                    if end != -1:
                        underline_text = text[i+2:end]
                        run = paragraph.add_run(underline_text)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.bold = bold
                        run.font.underline = True
                        i = end + 2
                        continue
                
                # Check for italic (*)
                elif text[i:i+1] == '*':
                    end = text.find('*', i + 1)
                    if end != -1:
                        italic_text = text[i+1:end]
                        run = paragraph.add_run(italic_text)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.bold = bold
                        run.font.italic = True
                        i = end + 1
                        continue
                
                # Regular text - accumulate until we hit a marker
                regular_end = i + 1
                while regular_end < len(text):
                    if text[regular_end:regular_end+3] == '///' or \
                       text[regular_end:regular_end+2] == '**' or \
                       text[regular_end:regular_end+2] == '__' or \
                       text[regular_end:regular_end+1] == '*':
                        break
                    regular_end += 1
                
                if regular_end > i:
                    regular_text = text[i:regular_end]
                    run = paragraph.add_run(regular_text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.bold = bold
                    if red:
                        run.font.color.rgb = RGBColor(178, 34, 34)
                    i = regular_end
                else:
                    i += 1
            
            # Remove any extra paragraphs
            while len(cell.paragraphs) > 1:
                cell._element.remove(cell.paragraphs[-1]._element)
        
        # Table 2: Job details (5 rows, 4 columns)
        # Index 1 because Table 1 is the header
        table2 = doc.tables[1]
        
        # Column 2 (index 1): Client, Product, Duration, Key No, Notes
        set_cell_text(table2.rows[0].cells[1], data.get('client', ''), bold=True)
        set_cell_text(table2.rows[1].cells[1], data.get('product', ''), bold=True)
        set_cell_text(table2.rows[2].cells[1], data.get('duration', ''), bold=True)
        set_cell_text(table2.rows[3].cells[1], data.get('keyNo', ''), bold=True, red=True)  # Red for Key No
        set_cell_text(table2.rows[4].cells[1], data.get('notes', ''), bold=True)
        
        # Column 4 (index 3): Contact, Ph No, Date, On Air, Prod $
        set_cell_text(table2.rows[0].cells[3], data.get('contact', ''), bold=True)
        set_cell_text(table2.rows[1].cells[3], data.get('phNo', ''), bold=True)
        set_cell_text(table2.rows[2].cells[3], data.get('date', ''), bold=True)
        set_cell_text(table2.rows[3].cells[3], data.get('onAir', ''), bold=True)
        set_cell_text(table2.rows[4].cells[3], data.get('prodCost', ''), bold=True)
        
        # Table 3: Script content (2 rows, 2 columns)
        # Index 2 because it's the third table
        table3 = doc.tables[2]
        
        # Row 2 (index 1): Video/Vision and Audio - use formatting function
        set_cell_text_with_formatting(table3.rows[1].cells[0], data.get('video', ''), bold=False)
        set_cell_text_with_formatting(table3.rows[1].cells[1], data.get('audio', ''), bold=False)
        
        # Table 4: Generator signature - NOT BOLD
        # Index 3 because it's the fourth table
        table4 = doc.tables[3]
        current_date = datetime.now().strftime('%d/%m/%y')
        set_cell_text(table4.rows[0].cells[0], f"Generated by AutoClapper 2.0 on {current_date}", bold=False)
        
        # Save the document with just the key number as filename
        key_no = data.get('keyNo', 'unknown').strip()
        output_filename = f"{key_no}.docx"
        output_path = os.path.join(output_dir, output_filename)
        doc.save(output_path)
        
        return output_path, None
        
    except Exception as e:
        return None, f"Error generating script: {str(e)}"


def generate_clapper_image(data, output_dir):
    """
    Generate the clapper image from template and user data.
    Places text at specified coordinates on the 1920x1080 clapper template.
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
        
        # Check if template exists
        if not os.path.exists(CLAPPER_TEMPLATE_PATH):
            return None, f"Clapper template not found at: {CLAPPER_TEMPLATE_PATH}"
        
        # Load the blank clapper template
        img = Image.open(CLAPPER_TEMPLATE_PATH)
        draw = ImageDraw.Draw(img)
        
        # Define text areas (top-left, bottom-right coordinates)
        text_areas = {
            'keyNo': {'coords': ((671, 420), (1498, 470)), 'text': data.get('keyNo', '')},
            'client': {'coords': ((671, 498), (1498, 548)), 'text': data.get('client', '')},
            'product': {'coords': ((671, 579), (1498, 629)), 'text': data.get('product', '')},
            'duration': {'coords': ((671, 659), (1498, 709)), 'text': data.get('duration', '')},
            'agency': {'coords': ((671, 738), (1498, 788)), 'text': data.get('agency', '')},
            'format': {'coords': ((671, 818), (1498, 868)), 'text': data.get('format', '16:9')}
        }
        
        # Try to load Arial font, fall back to default if not available
        try:
            # Try different possible Arial font paths on Windows
            font_paths = [
                'C:\\Windows\\Fonts\\arial.ttf',
                'C:\\Windows\\Fonts\\Arial.ttf',
                'arial.ttf'
            ]
            font = None
            for font_path in font_paths:
                try:
                    font = ImageFont.truetype(font_path, CLAPPER_FONT_SIZE)
                    break
                except:
                    continue
            
            if font is None:
                # Fall back to default font
                font = ImageFont.load_default()
        except:
            font = ImageFont.load_default()
        
        # Function to draw text in a given area
        def draw_text_in_box(draw, text, top_left, bottom_right, font):
            """Draw text left-aligned and vertically centered in the box."""
            if not text:
                return
            
            x1, y1 = top_left
            x2, y2 = bottom_right
            max_height = y2 - y1
            
            # Get text bounding box
            bbox = draw.textbbox((0, 0), text, font=font)
            text_height = bbox[3] - bbox[1]
            
            # Left-align horizontally, center vertically
            x = x1 + CLAPPER_OFFSET_X
            y = y1 + (max_height - text_height) // 2 + CLAPPER_OFFSET_Y
            
            # Draw the text in black
            draw.text((x, y), text, fill='black', font=font)
        
        # Draw text for each field
        for field_name, field_data in text_areas.items():
            top_left, bottom_right = field_data['coords']
            text = field_data['text']
            draw_text_in_box(draw, text, top_left, bottom_right, font)
        
        # Save the image
        key_no = data.get('keyNo', 'unknown').strip()
        output_filename = f"{key_no}.jpg"
        output_path = os.path.join(output_dir, output_filename)
        
        # Convert RGBA to RGB if necessary (JPEG doesn't support transparency)
        if img.mode == 'RGBA':
            rgb_img = Image.new('RGB', img.size, (255, 255, 255))
            rgb_img.paste(img, mask=img.split()[3])  # Use alpha channel as mask
            img = rgb_img
        
        img.save(output_path, 'JPEG', quality=80)
        
        return output_path, None
        
    except Exception as e:
        return None, f"Error generating clapper: {str(e)}"


@app.route('/')
def index():
    """Serve the HTML interface."""
    try:
        html_path = os.path.join(os.path.dirname(__file__), 'autoclapper-interface.html')
        with open(html_path, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return """
        <h1>AutoClapper 2.0</h1>
        <p>HTML interface file not found. Please open autoclapper-interface.html directly.</p>
        """, 404


@app.route('/generate', methods=['POST'])
def generate_documents():
    """
    Main endpoint that receives form data and generates both documents.
    """
    try:
        data = request.json
        output_dir = data.get('outputDirectory', '')
        
        # Validate output directory
        if not output_dir:
            return jsonify({
                'success': False,
                'error': 'No output directory specified'
            })
        
        if not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir)
            except Exception as e:
                return jsonify({
                    'success': False,
                    'error': f'Could not create output directory: {str(e)}'
                })
        
        # Generate script document
        script_path, script_error = generate_script_document(data, output_dir)
        if script_error:
            return jsonify({
                'success': False,
                'error': script_error
            })
        
        # Generate clapper image
        clapper_path, clapper_error = generate_clapper_image(data, output_dir)
        if clapper_error and "ready for integration" not in clapper_error:
            return jsonify({
                'success': False,
                'error': clapper_error
            })
        
        return jsonify({
            'success': True,
            'scriptPath': script_path,
            'clapperPath': clapper_path,
            'message': 'Documents generated successfully'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Server error: {str(e)}'
        })


@app.route('/health', methods=['GET'])
def health_check():
    """Simple health check endpoint."""
    template_status = {
        'script_template_exists': os.path.exists(SCRIPT_TEMPLATE_PATH),
        'clapper_template_exists': os.path.exists(CLAPPER_TEMPLATE_PATH),
        'script_template_path': SCRIPT_TEMPLATE_PATH,
        'clapper_template_path': CLAPPER_TEMPLATE_PATH
    }
    return jsonify(template_status)


@app.route('/transcribe', methods=['POST'])
def transcribe_audio():
    """Transcribe audio from uploaded file using Whisper."""
    try:
        # Check if Whisper is installed
        try:
            import whisper
        except ImportError:
            return jsonify({
                'success': False,
                'error': 'Whisper is not installed. Please run: pip install openai-whisper'
            })
        
        # Check if file was uploaded
        if 'audio' not in request.files:
            return jsonify({
                'success': False,
                'error': 'No audio file uploaded'
            })
        
        file = request.files['audio']
        
        if file.filename == '':
            return jsonify({
                'success': False,
                'error': 'No file selected'
            })
        
        # Save uploaded file temporarily
        import tempfile
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, f"autoclapper_temp_{file.filename}")
        file.save(temp_path)
        
        try:
            # Load Whisper model (using 'base' for good balance of speed/accuracy)
            # First time will download the model (~140MB)
            print("Loading Whisper model...")
            model = whisper.load_model("base")
            
            # Transcribe
            print(f"Transcribing {file.filename}...")
            result = model.transcribe(temp_path)
            transcription = result["text"].strip()
            
            print(f"Transcription complete: {len(transcription)} characters")
            
            return jsonify({
                'success': True,
                'transcription': transcription
            })
            
        finally:
            # Clean up temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Transcription error: {str(e)}'
        })


if __name__ == '__main__':
    try:
        # Create templates directory if it doesn't exist
        templates_dir = os.path.join(os.path.dirname(__file__), 'templates')
        if not os.path.exists(templates_dir):
            os.makedirs(templates_dir)
            print(f"Created templates directory at: {templates_dir}")
        
        print(f"Script template path: {SCRIPT_TEMPLATE_PATH}")
        print(f"Clapper template path: {CLAPPER_TEMPLATE_PATH}")
        print("Starting AutoClapper 2.0 server on http://localhost:5000")
        print("")
        print("NOTE: The first time you use audio transcription, Whisper will download a model file.")
        print("This is normal and only happens once.")
        print("")
        
        app.run(debug=True, port=5000)
    except Exception as e:
        print("")
        print("=" * 60)
        print("ERROR: Failed to start server")
        print("=" * 60)
        print(f"Error details: {e}")
        print("")
        print("This might be caused by:")
        print("  1. Port 5000 is already in use")
        print("  2. Missing dependencies")
        print("  3. Whisper installation issues")
        print("")
        print("Try closing other programs and running again.")
        print("=" * 60)
        input("Press Enter to exit...")

