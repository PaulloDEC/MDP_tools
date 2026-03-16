from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os
import tempfile
import traceback

app = Flask(__name__)
CORS(app)

# Output directory for transcriptions
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), 'transcriptions')
os.makedirs(OUTPUT_DIR, exist_ok=True)


def format_timestamp(seconds):
    """Convert seconds to SRT timestamp format (HH:MM:SS,mmm)"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    milliseconds = int((seconds % 1) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{milliseconds:03d}"


def create_srt(segments, output_path):
    """Create SRT subtitle file from Whisper segments"""
    with open(output_path, 'w', encoding='utf-8') as f:
        for i, segment in enumerate(segments, 1):
            start = format_timestamp(segment['start'])
            end = format_timestamp(segment['end'])
            text = segment['text'].strip()
            
            f.write(f"{i}\n")
            f.write(f"{start} --> {end}\n")
            f.write(f"{text}\n\n")


def create_txt(text, segments, output_path, include_timestamps):
    """Create plain text file"""
    with open(output_path, 'w', encoding='utf-8') as f:
        if include_timestamps:
            for segment in segments:
                timestamp = format_timestamp(segment['start'])
                f.write(f"[{timestamp}] {segment['text'].strip()}\n")
        else:
            f.write(text)


def create_docx(text, segments, output_path, include_timestamps, original_filename):
    """Create Word document with formatting"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Transcript: {original_filename}', 0)
    title.runs[0].font.color.rgb = None  # Default color
    
    # Add metadata
    metadata = doc.add_paragraph()
    metadata.add_run(f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n").font.size = Pt(9)
    metadata.add_run(f"Source: {original_filename}\n").font.size = Pt(9)
    metadata.add_run(f"Powered by OpenAI Whisper").font.size = Pt(9)
    
    doc.add_paragraph()  # Spacing
    
    # Add transcript
    if include_timestamps:
        for segment in segments:
            p = doc.add_paragraph()
            # Timestamp in bold
            timestamp_run = p.add_run(f"[{format_timestamp(segment['start'])}] ")
            timestamp_run.font.bold = True
            timestamp_run.font.size = Pt(11)
            # Text
            text_run = p.add_run(segment['text'].strip())
            text_run.font.size = Pt(11)
    else:
        # Add as continuous text with proper paragraphs
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                for run in p.runs:
                    run.font.size = Pt(11)
    
    doc.save(output_path)


@app.route('/')
def index():
    """Serve the HTML interface"""
    try:
        html_path = os.path.join(os.path.dirname(__file__), 'autotranscribe-interface.html')
        with open(html_path, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return """
        <h1>AutoTranscribe</h1>
        <p>HTML interface file not found. Please open autotranscribe-interface.html directly.</p>
        """, 404


@app.route('/transcribe_batch', methods=['POST'])
def transcribe_batch():
    """Transcribe a single file with specified options"""
    try:
        # Check if Whisper is installed
        try:
            import whisper
        except ImportError:
            return jsonify({
                'success': False,
                'error': 'Whisper is not installed. Please run: pip install openai-whisper'
            })
        
        # Get file
        if 'audio' not in request.files:
            return jsonify({'success': False, 'error': 'No audio file uploaded'})
        
        file = request.files['audio']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'})
        
        # Get settings
        model_size = request.form.get('model', 'base')
        language = request.form.get('language', 'auto')
        custom_vocab = request.form.get('custom_vocab', '')
        initial_prompt = request.form.get('initial_prompt', '')
        include_txt = request.form.get('include_txt', 'false') == 'true'
        include_docx = request.form.get('include_docx', 'false') == 'true'
        include_srt = request.form.get('include_srt', 'false') == 'true'
        include_timestamps = request.form.get('include_timestamps', 'false') == 'true'
        translate = request.form.get('translate', 'false') == 'true'
        
        # Save uploaded file temporarily
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, f"autotranscribe_temp_{file.filename}")
        file.save(temp_path)
        
        try:
            # Load Whisper model
            print(f"Loading Whisper model: {model_size}")
            model = whisper.load_model(model_size)
            
            # Prepare transcription options
            transcribe_options = {}
            
            # Handle language
            if language and language != 'auto':
                # Handle Australian English specially
                if language == 'en-AU':
                    transcribe_options['language'] = 'en'
                else:
                    transcribe_options['language'] = language
            
            # Build initial prompt with custom vocabulary
            prompt_parts = []
            if custom_vocab:
                # Add custom vocabulary to prompt
                vocab_list = [v.strip() for v in custom_vocab.split(',') if v.strip()]
                if vocab_list:
                    prompt_parts.append(f"Common terms: {', '.join(vocab_list)}.")
            
            if initial_prompt:
                prompt_parts.append(initial_prompt)
            
            # Add Australian English context if selected
            if language == 'en-AU':
                prompt_parts.append("Australian English with local place names and expressions.")
            
            if prompt_parts:
                transcribe_options['initial_prompt'] = ' '.join(prompt_parts)
                print(f"Using prompt: {transcribe_options['initial_prompt']}")
            
            # Enable word-level timestamps for confidence scores
            transcribe_options['word_timestamps'] = True
            
            # Transcribe or translate
            print(f"Processing {file.filename}...")
            if translate:
                result = model.transcribe(temp_path, task='translate', **transcribe_options)
            else:
                result = model.transcribe(temp_path, **transcribe_options)
            
            transcription = result["text"].strip()
            segments = result.get("segments", [])
            
            # Extract word-level data with confidence scores
            words_data = []
            for segment in segments:
                if 'words' in segment:
                    for word_info in segment['words']:
                        words_data.append({
                            'word': word_info.get('word', ''),
                            'start': word_info.get('start', 0),
                            'end': word_info.get('end', 0),
                            'probability': word_info.get('probability', 1.0)
                        })
            
            print(f"Transcription complete: {len(transcription)} characters, {len(words_data)} words")
            
            # Generate output files
            base_filename = os.path.splitext(file.filename)[0]
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_files = {}
            
            if include_txt:
                txt_filename = f"{base_filename}_{timestamp}.txt"
                txt_path = os.path.join(OUTPUT_DIR, txt_filename)
                create_txt(transcription, segments, txt_path, include_timestamps)
                output_files['txt_path'] = txt_filename
                print(f"Created TXT: {txt_path}")
            
            if include_docx:
                docx_filename = f"{base_filename}_{timestamp}.docx"
                docx_path = os.path.join(OUTPUT_DIR, docx_filename)
                create_docx(transcription, segments, docx_path, include_timestamps, file.filename)
                output_files['docx_path'] = docx_filename
                print(f"Created DOCX: {docx_path}")
            
            if include_srt:
                srt_filename = f"{base_filename}_{timestamp}.srt"
                srt_path = os.path.join(OUTPUT_DIR, srt_filename)
                create_srt(segments, srt_path)
                output_files['srt_path'] = srt_filename
                print(f"Created SRT: {srt_path}")
            
            return jsonify({
                'success': True,
                'transcription': transcription,
                'words': words_data,
                **output_files
            })
            
        except FileNotFoundError as e:
            return jsonify({
                'success': False,
                'error': 'FFmpeg is not installed. Please install FFmpeg and add it to your system PATH. See setup guide for instructions.'
            })
        finally:
            # Clean up temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
    except Exception as e:
        print(f"Transcription error: {traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error': f'Transcription error: {str(e)}'
        })


@app.route('/download', methods=['GET'])
def download_file():
    """Download a transcription file"""
    try:
        filename = request.args.get('path')
        if not filename:
            return "No filename specified", 400
        
        # Construct full path from filename
        file_path = os.path.join(OUTPUT_DIR, filename)
        
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return "File not found", 404
        
        return send_file(file_path, as_attachment=True)
    except Exception as e:
        print(f"Download error: {str(e)}")
        return f"Download error: {str(e)}", 500


@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    try:
        import whisper
        whisper_installed = True
    except ImportError:
        whisper_installed = False
    
    # Check for FFmpeg
    import subprocess
    try:
        subprocess.run(['ffmpeg', '-version'], capture_output=True, check=True)
        ffmpeg_installed = True
    except (subprocess.CalledProcessError, FileNotFoundError):
        ffmpeg_installed = False
    
    return jsonify({
        'status': 'ok',
        'whisper_installed': whisper_installed,
        'ffmpeg_installed': ffmpeg_installed,
        'output_directory': OUTPUT_DIR
    })


if __name__ == '__main__':
    print("=" * 60)
    print("AutoTranscribe - Batch Audio Transcription")
    print("=" * 60)
    print(f"Output directory: {OUTPUT_DIR}")
    print("Starting server on http://localhost:5001")
    print("")
    print("Note: First time using a model will download it (~140MB-2.9GB)")
    print("=" * 60)
    print("")
    
    app.run(debug=True, port=5001)
